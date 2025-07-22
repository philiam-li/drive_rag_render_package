

import os
import io
import json
import fitz  # PyMuPDF
import docx
import requests
from datetime import datetime, timedelta, timezone
from fastapi import FastAPI, Request
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from starlette.responses import Response
from linebot import LineBotApi
from linebot.models import TextSendMessage
import google.generativeai as genai
from dateutil.parser import parse as parse_date

# ✅ FastAPI 應用初始化
app = FastAPI()

# ✅ 環境變數
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
SERVICE_ACCOUNT_INFO = json.loads(os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON"))
LINE_CHANNEL_ACCESS_TOKEN = os.getenv("LINE_CHANNEL_ACCESS_TOKEN")
LINE_USER_ID = os.getenv("LINE_USER_ID")

INPUT_FOLDER_NAME = "spec-inbox"
OUTPUT_FOLDER_NAME = "spec-outbox"
TEMP_FOLDER = "temp"

# ✅ 初始化 Gemini & LINE & Google Drive
genai.configure(api_key=GEMINI_API_KEY)
gemini = genai.GenerativeModel("gemini-2.0-flash")
line_bot_api = LineBotApi(LINE_CHANNEL_ACCESS_TOKEN)

SCOPES = ["https://www.googleapis.com/auth/drive"]
credentials = service_account.Credentials.from_service_account_info(SERVICE_ACCOUNT_INFO, scopes=SCOPES)
drive_service = build("drive", "v3", credentials=credentials)

# ✅ LINE 訊息通知
def send_line_message_to_self(message: str):
    try:
        line_bot_api.push_message(LINE_USER_ID, TextSendMessage(text=message))
        print("✅ LINE 訊息已送出")
    except Exception as e:
        print("❌ LINE 訊息發送失敗：", e)

# ✅ Google Drive 工具
def get_folder_id_by_name(folder_name):
    results = drive_service.files().list(
        q=f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder'",
        fields="files(id, name)"
    ).execute()
    items = results.get("files", [])
    return items[0]["id"] if items else None

def list_new_files_in_folder(folder_id, minutes=3):
    cutoff = datetime.now(timezone.utc) - timedelta(minutes=minutes)
    results = drive_service.files().list(
        q=f"'{folder_id}' in parents and (mimeType='application/pdf' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or mimeType='text/plain')",
        fields="files(id, name, mimeType, createdTime)"
    ).execute()
    return [
        f for f in results.get("files", [])
        if parse_date(f["createdTime"]) > cutoff
    ]

def list_files_in_folder(folder_id):
    results = drive_service.files().list(
        q=f"'{folder_id}' in parents",
        fields="files(id, name, mimeType)"
    ).execute()
    return results.get("files", [])

def save_file_to_temp(file_id, file_name):
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    os.makedirs(TEMP_FOLDER, exist_ok=True)
    local_path = os.path.join(TEMP_FOLDER, file_name)
    with open(local_path, "wb") as f:
        f.write(fh.getbuffer())
    return local_path

# ✅ 分析邏輯
def extract_text_from_file(file_path):
    if file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif file_path.endswith(".pdf"):
        try:
            doc = fitz.open(file_path)
            return "\n".join([page.get_text() for page in doc])
        except Exception as e:
            print(f"❌ 無法解析 PDF：{e}")
            return ""
    print("⚠️ 不支援的檔案格式：", file_path)
    return ""

def analyze_text_with_gemini(text, chunk_size=8000):
    if not text.strip():
        return "⚠️ 無法從文件中擷取任何內容，請確認格式或重新上傳。"

    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    print(f"📚 總共分成 {len(chunks)} 段")

    full_response = []
    for i, chunk in enumerate(chunks, 1):
        prompt = f"""你是一位資深系統分析師，請針對以下文件片段進行風險審查：

--- 第 {i} 段 ---
{chunk}
---

請條列三類事項（繁體中文）：
1. 🔺 潛在風險
2. 💡 建議調整
3. ❓ 需補充釐清
"""
        print(f"🧠 分析第 {i} 段...")
        try:
            result = gemini.generate_content(prompt)
            full_response.append(f"📍 第 {i} 段分析：\n{result.text.strip()}\n")
        except Exception as e:
            full_response.append(f"❌ 第 {i} 段錯誤：{e}\n")

    return "\n\n".join(full_response)

# ✅ 輸出報告
def write_summary_to_docx(summary, output_path, original_filename=None):
    doc = Document()
    doc.add_heading('系統規格風險分析報告', level=1)
    if original_filename:
        doc.add_paragraph(f"📄 來源檔案：{original_filename}")
    doc.add_heading("📌 Gemini 分析結果", level=2)
    for line in summary.splitlines():
        doc.add_paragraph(line, style='List Bullet')
    doc.save(output_path)
    
def upload_file_to_folder(folder_id, local_path, file_name):
    file_metadata = {"name": file_name, "parents": [folder_id]}
    media = MediaFileUpload(local_path, resumable=False)  # 改成 simple upload
    drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields="id",
        supportsAllDrives=False  # 明確說這不是共用雲端硬碟
    ).execute()

# ✅ webhook 主處理邏輯
@app.post("/webhook")
def run_agent():
    input_folder_id = get_folder_id_by_name(INPUT_FOLDER_NAME)
    output_folder_id = get_folder_id_by_name(OUTPUT_FOLDER_NAME)
    if not input_folder_id or not output_folder_id:
        return {"error": "❌ 找不到 Google Drive 資料夾"}

    files = list_new_files_in_folder(input_folder_id, minutes=3)
    existing_outputs = [f["name"] for f in list_files_in_folder(output_folder_id)]
    os.makedirs(TEMP_FOLDER, exist_ok=True)

    for file in files:
        file_name = file["name"]
        if file_name.startswith("報告_") or not file_name.endswith((".pdf", ".docx", ".txt")):
            print(f"⏩ 跳過檔案：{file_name}")
            continue

        output_name = f"報告_{file_name}.docx"
        if output_name in existing_outputs:
            print(f"⏩ 已存在報告，跳過：{output_name}")
            continue

        print(f"📥 處理中：{file_name}")
        local_path = save_file_to_temp(file["id"], file_name)
        text = extract_text_from_file(local_path)
        summary = analyze_text_with_gemini(text)
        docx_path = os.path.join(TEMP_FOLDER, output_name)
        write_summary_to_docx(summary, docx_path, original_filename=file_name)
        upload_file_to_folder(output_folder_id, docx_path, output_name)
        print(f"✅ 已完成並上傳報告：{output_name}")
        send_line_message_to_self(f"✅ 分析完成：{output_name}")

    return {"message": "🎉 所有檔案處理完成"}
